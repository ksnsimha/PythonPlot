#!/usr/bin/env python
# coding: utf-8

# In[28]:


def downloadChart():
    import pandas as pd
    
    csvFile = input("Enter the full filepath of the csv: ")
    
    
    df = pd.read_csv(csvFile,error_bad_lines=False)
    import matplotlib.pyplot as plt
    X = list(df.iloc[:,0])
    Y = list(df.iloc[:,1])
    plt.bar(X,Y,color='g')
    plt.title("Sample Bar Graph")
    plt.xlabel("X-AXIS")
    plt.ylabel("Y-AXIS")
    plt.savefig("TEMP.PNG")
def exportPnGToDoc():
    from docx import Document
    from docx.shared import Inches
    from docx import Document
    from docx.shared import Inches
    import os
    docName = input("Enter the full filepath of the word document: ")
   
    
    if not os.path.isfile(docName):
        document = Document()
        document.save(docName)
        print("New document created in the path "+docName)
    else:
        document = Document(docName)
        print("Adding the barchart to the existing document "+docName)
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('Sample Paragraph')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    
    document.add_picture('TEMP.PNG', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
 

    document.add_page_break()
    

    document.save(docName)
    
    if os.path.exists('TEMP.PNG'):
        os.remove('TEMP.PNG')
downloadChart()
exportPnGToDoc()
